"""Word 文档图片插入工具，负责按版面约束写入本地图片。"""

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt
from PIL import Image


def _insert_picture_with_constraints(
    document: DocxDocument,
    image_path: str,
    max_width_cm: float = 15.5,
    max_height_cm: float = 9.5,
) -> None:
    """
    插入图片并限制最大宽高，避免图片过高导致版面被大面积占用。
    保持原始纵横比，不做拉伸。
    """
    with Image.open(image_path) as image:
        width_px, height_px = image.size

    if width_px <= 0 or height_px <= 0:
        document.add_picture(image_path, width=Cm(max_width_cm))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    width_cm = max_width_cm
    height_cm = width_cm * height_px / width_px
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * width_px / height_px

    document.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))
    p = document.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 强制单倍行距，避免被 Normal 样式的固定值（22pt）截断导致图片被文本覆盖
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Pt(0)
