import datetime
from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from services.thesis.document.formatting import _apply_fixed_line_spacing, _set_run_font, _toc_int
from services.thesis.document.toc import _add_bookmark


def _add_cover_page(
    document: DocxDocument,
    title: str,
    author: str,
    advisor: str,
    degree_type: str,
    major: str,
    school: str,
    year_month: str,
) -> None:
    """构建封面页。"""

    def _blank(lines: int) -> None:
        for _ in range(lines):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    _blank(6)

    p_main = document.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main.paragraph_format.first_line_indent = Pt(0)
    # Cover title uses a much larger font than body text; if it inherits the
    # Normal style's fixed 22pt line height, WPS clips the glyphs vertically.
    p_main.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run_main = p_main.add_run("学士学位论文")
    _set_run_font(run_main, zh_font="黑体", size_pt=36, bold=True)
    p_main.paragraph_format.space_after = Pt(24)

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run_title = p_title.add_run(title)
    _set_run_font(run_title, size_pt=18, underline=True)
    p_title.paragraph_format.space_after = Pt(36)

    fields = [
        ("作者姓名", author),
        ("指导教师", advisor),
        ("学位类别", degree_type),
        ("专    业", major),
        ("学院（系）", school),
    ]
    for label, value in fields:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(3.5)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        run_label = paragraph.add_run(f"{label}：")
        _set_run_font(run_label, size_pt=14)
        run_value = paragraph.add_run(value)
        _set_run_font(run_value, size_pt=14, bold=True)

    _blank(8)

    if not year_month:
        now = datetime.datetime.now()
        year_month = f"{now.year} 年 {now.month} 月"

    p_date = document.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.first_line_indent = Pt(0)
    run_date = p_date.add_run(year_month)
    _set_run_font(run_date, size_pt=14, bold=True)

    document.add_page_break()


def _apply_full_border(table: Any) -> None:
    """给承诺书/授权书信息表添加全边框。"""
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for old in tc_pr.findall(qn("w:tcBorders")):
                tc_pr.remove(old)
            borders = OxmlElement("w:tcBorders")
            for name in ("top", "bottom", "left", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), "000000")
                borders.append(border)
            tc_pr.append(borders)


def _add_signature_page(
    document: DocxDocument,
    page_title: str,
    body_title: str,
    body_text: str,
    title: str,
    author: str,
    advisor: str,
    major: str,
    school: str,
    student_id: str,
    student_class: str,
) -> None:
    """添加诚信承诺书或版权使用授权书。"""
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(18)
    run_title = p_title.add_run(page_title)
    _set_run_font(run_title, zh_font="黑体", size_pt=16, bold=True)

    table = document.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("毕业论文（设计）题目", title, "学生姓名", author),
        ("学院（系）", school, "专业", major),
        ("班级", student_class, "学号", student_id),
        ("指导教师", advisor, "实践导师", ""),
    ]
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _set_run_font(run, size_pt=10.5)
    _apply_full_border(table)

    document.add_paragraph()
    p_body_title = document.add_paragraph()
    p_body_title.paragraph_format.first_line_indent = Pt(0)
    run_body_title = p_body_title.add_run(body_title)
    _set_run_font(run_body_title, zh_font="黑体", size_pt=12, bold=True)

    for line in body_text.split("\n"):
        text = line.strip()
        if not text:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        run = paragraph.add_run(text)
        _set_run_font(run, size_pt=12)

    p_sign = document.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.paragraph_format.first_line_indent = Pt(0)
    run_sign = p_sign.add_run("学生签名：              年    月    日")
    _set_run_font(run_sign, size_pt=12)
    document.add_page_break()


def _add_integrity_page(
    document: DocxDocument,
    title: str,
    author: str,
    advisor: str,
    major: str,
    school: str,
    student_id: str,
    student_class: str,
) -> None:
    body = (
        "本人慎重承诺和声明：\n"
        "我承诺在毕业论文（设计）活动中遵守学校有关规定，恪守学术规范，"
        "在本人的毕业论文中未剽窃、抄袭他人的学术观点、思想和成果，"
        "未篡改研究数据，如有违规行为发生，我愿承担一切责任，接受学校的处理。"
    )
    _add_signature_page(
        document,
        "本科生毕业论文（设计）诚信承诺书",
        "诚信承诺",
        body,
        title,
        author,
        advisor,
        major,
        school,
        student_id,
        student_class,
    )


def _add_copyright_page(
    document: DocxDocument,
    title: str,
    author: str,
    advisor: str,
    major: str,
    school: str,
    student_id: str,
    student_class: str,
) -> None:
    body = (
        "本毕业论文（设计）是本人在校期间所完成学业的组成部分，是在学校教师的指导下完成的。"
        "因此，本人特授权学校可将本毕业论文（设计）的全部或部分内容编入有关书籍、数据库保存，"
        "可采用复制、印刷、网页制作等方式将论文（设计）文本和经过编辑、批注等处理的论文（设计）"
        "文本提供给读者查阅、参考，可向有关学术部门和国家有关教育主管部门呈送复印件和电子文档。"
        "本毕业论文（设计）无论做何种处理，必须尊重本人的著作权，署明本人姓名。"
    )
    _add_signature_page(
        document,
        "本科生毕业论文（设计）版权使用授权书",
        "使用授权",
        body,
        title,
        author,
        advisor,
        major,
        school,
        student_id,
        student_class,
    )


def _add_abstract_zh_page(document: DocxDocument, abstract: str, keywords: str) -> None:
    """中文摘要页。"""
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("摘    要")
    _set_run_font(run_title, zh_font="黑体", size_pt=16, bold=True)

    if abstract:
        for para_text in abstract.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            run = paragraph.add_run(para_text)
            _set_run_font(run, size_pt=12)

    if keywords:
        p_kw = document.add_paragraph()
        p_kw.paragraph_format.first_line_indent = Pt(0)
        p_kw.paragraph_format.space_before = Pt(12)
        run_label = p_kw.add_run("【关键词】")
        _set_run_font(run_label, zh_font="黑体", size_pt=12, bold=True)
        run_kw = p_kw.add_run(keywords)
        _set_run_font(run_kw, size_pt=12)

    document.add_page_break()


def _add_abstract_en_page(document: DocxDocument, abstract: str, keywords: str) -> None:
    """英文摘要页。"""
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("ABSTRACT")
    _set_run_font(
        run_title,
        zh_font="Times New Roman",
        en_font="Times New Roman",
        size_pt=16,
        bold=True,
    )

    if abstract:
        for para_text in abstract.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(48)
            paragraph.paragraph_format.space_before = Pt(6)
            run = paragraph.add_run(para_text)
            _set_run_font(
                run,
                zh_font="Times New Roman",
                en_font="Times New Roman",
                size_pt=12,
            )

    if keywords:
        p_kw = document.add_paragraph()
        p_kw.paragraph_format.first_line_indent = Pt(0)
        p_kw.paragraph_format.space_before = Pt(12)
        run_label = p_kw.add_run("【KEY WORDS】")
        _set_run_font(
            run_label,
            zh_font="Times New Roman",
            en_font="Times New Roman",
            size_pt=12,
            bold=True,
        )
        run_kw = p_kw.add_run(keywords)
        _set_run_font(
            run_kw,
            zh_font="Times New Roman",
            en_font="Times New Roman",
            size_pt=12,
        )



def _add_acknowledgment_page(
    document: DocxDocument,
    acknowledgment: str,
    bookmark: dict[str, object] | None = None,
) -> None:
    """致谢页。"""
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("致    谢")
    _set_run_font(run_title, zh_font="黑体", size_pt=16, bold=True)
    if bookmark:
        _add_bookmark(p_title, str(bookmark["bookmark"]), _toc_int(bookmark["bookmark_id"]))

    if not acknowledgment:
        return

    for para_text in acknowledgment.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(24)
        run = paragraph.add_run(para_text)
        _set_run_font(run, size_pt=12)


def _add_references_page(
    document: DocxDocument,
    references: str,
    bookmark: dict[str, object] | None = None,
) -> None:
    """参考文献页。"""
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("参考文献")
    _set_run_font(run_title, zh_font="黑体", size_pt=16, bold=True)
    if bookmark:
        _add_bookmark(p_title, str(bookmark["bookmark"]), _toc_int(bookmark["bookmark_id"]))

    if not references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        run = paragraph.add_run("（参考文献未生成，可能因未配置检索服务或检索结果不足）")
        _set_run_font(run, size_pt=10.5)
        return

    for line in references.splitlines():
        line = line.strip()
        if not line:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.8)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        _apply_fixed_line_spacing(paragraph.paragraph_format, pt=18)
        run = paragraph.add_run(line)
        _set_run_font(run, size_pt=10.5)
