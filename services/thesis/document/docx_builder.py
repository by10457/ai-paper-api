"""论文 Word 文档构建入口，负责组织封面、摘要、正文、图表、致谢和参考文献。"""

import datetime
import re
from pathlib import Path
from typing import Any

from docx import Document as DocumentFactory
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from services.thesis.document.figures import _insert_picture_with_constraints
from services.thesis.document.formatting import _set_run_font, _toc_int
from services.thesis.document.inline import _add_markdown_text_to_paragraph, _add_table, _collect_table_lines
from services.thesis.document.pages import (
    _add_abstract_en_page,
    _add_abstract_zh_page,
    _add_acknowledgment_page,
    _add_copyright_page,
    _add_cover_page,
    _add_integrity_page,
    _add_references_page,
)
from services.thesis.document.sections import (
    _make_blank_section,
    _setup_body_section,
    _setup_front_matter_section,
)
from services.thesis.document.styles import _init_styles, _setup_page
from services.thesis.document.toc import FIGURE_BLOCK_PATTERN, _add_bookmark, _add_toc_page, _pre_scan_headings


def build_word_document(
    full_text: str,
    placeholders: list[dict[str, Any]],
    image_paths: dict[int, str | None],
    output_path: str = "public/output/thesis/thesis.docx",
    title: str = "论文题目",
    author: str = "作者姓名",
    advisor: str = "指导教师",
    degree_type: str = "学士",
    major: str = "专业名称",
    school: str = "XX大学XX学院",
    year_month: str = "",
    student_id: str = "",
    student_class: str = "",
    abstract_zh: str = "",
    abstract_en: str = "",
    keywords_zh: str = "",
    keywords_en: str = "",
    acknowledgment: str = "",
    references: str = "",
) -> str:
    """将论文正文、图片与前后置页面构造成 Word 文档。"""
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document = DocumentFactory()

    _setup_page(document)
    _init_styles(document)

    # Section 1：封面 + 诚信承诺书 + 版权使用授权书；不输出页码。
    _add_cover_page(
        document,
        title,
        author,
        advisor,
        degree_type,
        major,
        school,
        year_month,
    )
    _add_integrity_page(document, title, author, advisor, major, school, student_id, student_class)
    _add_copyright_page(document, title, author, advisor, major, school, student_id, student_class)

    # Section 2：中文摘要，从罗马页码 I 开始。
    zh_section = _make_blank_section(document, WD_SECTION.CONTINUOUS)
    _setup_front_matter_section(zh_section, start=1)
    _add_abstract_zh_page(document, abstract_zh, keywords_zh)

    # Section 3：英文摘要，罗马页码自然延续。
    en_section = _make_blank_section(document, WD_SECTION.CONTINUOUS)
    _setup_front_matter_section(en_section)
    _add_abstract_en_page(document, abstract_en, keywords_en)

    # 预扫描正文标题，生成目录条目列表
    toc_entries = _pre_scan_headings(full_text, title=title)

    # Section 4：目录（可见条目 + PAGEREF 动态页码），罗马页码自然延续。
    toc_section = _make_blank_section(document)
    _setup_front_matter_section(toc_section)
    _add_toc_page(document, toc_entries, full_text=full_text)

    # Section 5：正文 + 参考文献 + 致谢，阿拉伯页码从 1 起。
    _make_blank_section(document)
    _setup_body_section(document, title)

    segments = re.split(FIGURE_BLOCK_PATTERN, full_text, flags=re.DOTALL)
    placeholder_idx = 0

    # 目录条目迭代器，用于在正文标题上打 bookmark
    _toc_iter = iter(toc_entries)
    _next_toc = next(_toc_iter, None)

    for segment in segments:
        lines = [line.strip() for line in segment.strip().split("\n")]
        i = 0
        while i < len(lines):
            line = lines[i]

            if not line:
                i += 1
                continue

            if line == "---pagebreak---":
                document.add_page_break()
                i += 1
                continue

            if line.startswith("|") and "|" in line[1:]:
                rows, next_i = _collect_table_lines(lines, i)
                if rows:
                    _add_table(document, rows)
                i = next_i
                continue

            if re.match(r"^表\s?\d+(?:[-.]\d+)?\s", line):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                _add_markdown_text_to_paragraph(paragraph, line)
                for run in paragraph.runs:
                    _set_run_font(run, size_pt=10.5, bold=True)
                i += 1
                continue

            if line.startswith("### "):
                heading_text = line[4:]
                paragraph = document.add_heading(heading_text, level=3)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if _next_toc and _next_toc["text"] == heading_text.strip() and _next_toc["level"] == 3:
                    _add_bookmark(paragraph, str(_next_toc["bookmark"]), _toc_int(_next_toc["bookmark_id"]))
                    _next_toc = next(_toc_iter, None)
            elif line.startswith("## "):
                heading_text = line[3:]
                paragraph = document.add_heading(heading_text, level=2)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if _next_toc and _next_toc["text"] == heading_text.strip() and _next_toc["level"] == 2:
                    _add_bookmark(paragraph, str(_next_toc["bookmark"]), _toc_int(_next_toc["bookmark_id"]))
                    _next_toc = next(_toc_iter, None)
            elif line.startswith("# "):
                heading_text = line[2:]
                paragraph = document.add_heading(heading_text, level=1)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if _next_toc and _next_toc["text"] == heading_text.strip() and _next_toc["level"] == 1:
                    _add_bookmark(paragraph, str(_next_toc["bookmark"]), _toc_int(_next_toc["bookmark_id"]))
                    _next_toc = next(_toc_iter, None)
            elif line.startswith("```"):
                pass
            elif line.startswith("- "):
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.first_line_indent = Pt(0)
                _add_markdown_text_to_paragraph(paragraph, line[2:])
            elif re.match(r"^\d+\.\s+", line):
                paragraph = document.add_paragraph(style="List Number")
                paragraph.paragraph_format.first_line_indent = Pt(0)
                text = re.sub(r"^\d+\.\s+", "", line)
                _add_markdown_text_to_paragraph(paragraph, text)
            else:
                paragraph = document.add_paragraph()
                _add_markdown_text_to_paragraph(paragraph, line)

            i += 1

        if placeholder_idx >= len(placeholders):
            continue

        placeholder = placeholders[placeholder_idx]
        image_path = image_paths.get(placeholder["index"])
        has_image = bool(image_path and Path(image_path).exists())

        if has_image and image_path:
            _insert_picture_with_constraints(document, image_path)

        caption = placeholder.get("caption", f"图{placeholder_idx + 1}")
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.first_line_indent = Pt(0)
        caption_run = caption_paragraph.add_run(caption)
        _set_run_font(caption_run, size_pt=10.5, bold=True)

        placeholder_idx += 1

    ref_bookmark = next((entry for entry in toc_entries if entry["text"] == "参考文献"), None)
    ack_bookmark = next((entry for entry in toc_entries if entry["text"] == "致谢"), None)

    document.add_page_break()
    _add_references_page(document, references, bookmark=ref_bookmark)

    document.add_page_break()
    _add_acknowledgment_page(document, acknowledgment, bookmark=ack_bookmark)

    # 修正文档核心属性，让 Windows 资源管理器正确识别并显示 Word 图标。
    # python-docx 默认模板的创建时间是 2013 年，creator 是 "python-docx"，
    # 这会导致 Windows 无法正确生成预览图标。
    now = datetime.datetime.now()
    core = document.core_properties
    core.title = title
    core.author = author
    core.created = now
    core.modified = now
    core.last_modified_by = author
    core.revision = 1
    core.description = ""

    document.save(output_path)
    return output_path
