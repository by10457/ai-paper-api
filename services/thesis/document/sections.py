from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.thesis.document.formatting import _set_page_numbering, _set_run_font


def _clear_header_footer(section: Any) -> None:
    """清空当前 section 的页眉页脚。"""
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for paragraph in section.header.paragraphs:
        paragraph.clear()
    for paragraph in section.footer.paragraphs:
        paragraph.clear()


def _copy_page_layout_from_first(document: DocxDocument, section: Any) -> None:
    """让新增 section 继承第一页的纸张和页边距设置。"""
    first = document.sections[0]
    section.page_width = first.page_width
    section.page_height = first.page_height
    section.top_margin = first.top_margin
    section.bottom_margin = first.bottom_margin
    section.left_margin = first.left_margin
    section.right_margin = first.right_margin
    section.header_distance = first.header_distance
    section.footer_distance = first.footer_distance


def _make_blank_section(document: DocxDocument, section_type: Any = WD_SECTION.NEW_PAGE) -> Any:
    """新增一个空白 section，不带页眉页脚。"""
    section = document.add_section(section_type)
    _copy_page_layout_from_first(document, section)
    _clear_header_footer(section)
    return section


def _add_page_number_footer(section: Any, cached_text: str = "1") -> None:
    """在页脚居中加入 PAGE 字段。"""
    footer = section.footer
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.clear()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_f = p_footer.add_run()
    r_element = run_f._element

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_element.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r_element.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_element.append(fld_sep)

    num_run = p_footer.add_run(cached_text)
    _set_run_font(
        num_run,
        zh_font="Times New Roman",
        en_font="Times New Roman",
        size_pt=10.5,
    )

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p_footer.add_run()._element.append(fld_end)


def _setup_front_matter_section(section: Any, start: int | None = None) -> None:
    """配置摘要/目录 section：罗马页码，start=None 时自然延续。"""
    _clear_header_footer(section)
    _set_page_numbering(section, start=start, number_format="upperRoman")
    _add_page_number_footer(section, cached_text="Ⅰ" if start == 1 else "")


def _setup_body_section(document: DocxDocument, title: str) -> None:
    """配置正文 section 的页眉和页脚页码。"""
    section = document.sections[-1]
    _copy_page_layout_from_first(document, section)
    _clear_header_footer(section)
    _set_page_numbering(section, start=1)

    header = section.header
    p_header = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p_header.clear()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = p_header.add_run(title)
    _set_run_font(run_h, zh_font="Times New Roman", en_font="Times New Roman", size_pt=10.5)
    _add_page_number_footer(section, cached_text="1")
