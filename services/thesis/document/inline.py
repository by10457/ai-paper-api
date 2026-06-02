import re
from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from services.thesis.document.formatting import _set_run_font

CITATION_PATTERN = re.compile(r"\[(\d{1,3})\]")


def _is_table_separator(line: str) -> bool:
    """判断是否为 Markdown 表格分隔行。"""
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.match(r"^[-:]+$", c) for c in cells) if cells else False


def _parse_table_line(line: str) -> list[str]:
    """解析单行 Markdown 表格。"""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _collect_table_lines(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """收集连续的 Markdown 表格行。"""
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and "|" in lines[i]:
        if _is_table_separator(lines[i]):
            i += 1
            continue
        rows.append(_parse_table_line(lines[i]))
        i += 1
    return rows, i


def _add_markdown_text_to_paragraph(
    paragraph: Any,
    text: str,
    is_header: bool = False,
    is_table: bool = False,
) -> None:
    """解析 Markdown 内联加粗并写入段落。"""
    text = text.replace("~", "")
    parts = re.split(r"(\*\*.*?\*\*|\[\d{1,3}\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif CITATION_PATTERN.fullmatch(part):
            run = paragraph.add_run(part)
            run.font.superscript = True
        else:
            run = paragraph.add_run(part)
            if is_header:
                run.bold = True
        _set_run_font(run, zh_font="宋体")
        if is_table:
            run.font.size = Pt(10.5)


def _add_table(document: DocxDocument, rows: list[list[str]]) -> None:
    """将二维数据写入 Word 三线表；第一行视为表头。"""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                continue
            cell = table.rows[i].cells[j]
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_markdown_text_to_paragraph(
                paragraph,
                cell_text,
                is_header=(i == 0),
                is_table=True,
            )
            for run in paragraph.runs:
                _set_run_font(run, size_pt=10.5, bold=True if i == 0 else None)

    _apply_three_line_table(table)


def _apply_three_line_table(table: Any) -> None:
    """三线表：首行顶线/表头底线、末行底线，左右开放。"""
    row_count = len(table.rows)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for old in tc_pr.findall(qn("w:tcBorders")):
                tc_pr.remove(old)

            borders = OxmlElement("w:tcBorders")
            bottom_sz = "12" if row_idx == row_count - 1 else ("4" if row_idx == 0 else "nil")
            for name, sz in (
                ("top", "12" if row_idx == 0 else "nil"),
                ("bottom", bottom_sz),
                ("left", "nil"),
                ("right", "nil"),
                ("insideH", "nil"),
                ("insideV", "nil"),
            ):
                border = OxmlElement(f"w:{name}")
                if sz == "nil":
                    border.set(qn("w:val"), "nil")
                else:
                    border.set(qn("w:val"), "single")
                    border.set(qn("w:sz"), sz)
                    border.set(qn("w:color"), "000000")
                borders.append(border)
            tc_pr.append(borders)
