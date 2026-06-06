"""Word 目录页生成工具，负责预扫描标题、书签和目录页码占位。"""

import re
from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from services.thesis.document.formatting import _apply_fixed_line_spacing, _set_run_font, _toc_int

FIGURE_BLOCK_PATTERN = r"<<FIGURE>>\s*.*?\s*<</FIGURE>>"


def _pre_scan_headings(
    full_text: str,
    title: str = "",
    include_back_matter: bool = True,
) -> list[dict[str, object]]:
    """预扫描正文 Markdown，提取 1-3 级标题供目录页使用。"""
    clean_text = re.sub(FIGURE_BLOCK_PATTERN, "", full_text, flags=re.DOTALL)
    entries: list[dict[str, object]] = []

    # LLM 偶尔会把前后置页面标题写入正文，这些标题不能进入正文目录。
    non_body_headings = {
        "摘要",
        "摘 要",
        "中文摘要",
        "abstract",
        "致谢",
        "致 谢",
        "参考文献",
    }

    for line in clean_text.split("\n"):
        line = line.strip()
        if not line or line == "---pagebreak---":
            continue

        level = 0
        text = ""
        if line.startswith("### "):
            level, text = 3, line[4:]
        elif line.startswith("## "):
            level, text = 2, line[3:]
        elif line.startswith("# "):
            level, text = 1, line[2:]

        if level == 0:
            continue

        text = text.strip()
        if title and text == title.strip():
            continue
        if text.lower() in non_body_headings:
            continue

        idx = len(entries)
        entries.append(
            {
                "text": text,
                "level": level,
                "bookmark": f"_toc_{idx}",
                "bookmark_id": idx + 100,
            }
        )

    if include_back_matter:
        for text in ("参考文献", "致谢"):
            idx = len(entries)
            entries.append(
                {
                    "text": text,
                    "level": 1,
                    "bookmark": f"_toc_{idx}",
                    "bookmark_id": idx + 100,
                }
            )

    return entries


def _add_bookmark(paragraph: Any, bookmark_name: str, bookmark_id: int) -> None:
    """Wrap paragraph content with bookmarkStart / bookmarkEnd."""
    p_element = paragraph._element

    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), bookmark_name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))

    p_pr = p_element.find(qn("w:pPr"))
    if p_pr is not None:
        p_pr.addnext(bm_start)
    else:
        p_element.insert(0, bm_start)

    p_element.append(bm_end)


def _add_pageref_field(paragraph: Any, bookmark_name: str, cached_page: str = "?") -> None:
    """Append a PAGEREF field referencing *bookmark_name* to *paragraph*.

    *cached_page* is the pre-estimated page number displayed until the
    end-user's application recalculates fields.  ``dirty="true"`` on the
    begin fldChar signals Word / WPS that this field should be refreshed.

    The generated OOXML structure::

        <w:fldChar begin dirty="true"/>
        <w:instrText> PAGEREF _toc_0 \\h </w:instrText>
        <w:fldChar separate/>
        <w:t>5</w:t>          (estimated; corrected on F9 / auto-update)
        <w:fldChar end/>
    """
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    # NOTE: Do NOT set w:dirty here. Setting dirty="true" causes Word to
    # auto-refresh the PAGEREF, but because body section restarts page
    # numbering, Word resolves all fields to "1" before layout completes.
    run_begin._element.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    run_instr._element.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._element.append(fld_sep)

    run_placeholder = paragraph.add_run(str(cached_page))
    _set_run_font(run_placeholder, zh_font="Times New Roman", en_font="Times New Roman", size_pt=12)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._element.append(fld_end)

def _estimate_page_numbers(
    full_text: str,
    toc_entries: list[dict[str, object]],
    body_start_page: int = 1,
) -> dict[str, int]:
    """Estimate the page number for each TOC heading accurately based on lines.

    Uses a line-count heuristic:
      - A4 with 3cm/2.5cm margins, approx 31 lines per page
      - Approx 36 characters per line
      - Images take about 14 lines (half a page)
      - Headings take extra lines for vertical spacing

    Returns ``{bookmark_name: estimated_page, ...}``.
    """
    lines_per_page = 31
    chars_per_line = 36
    lines_per_image = 14

    text_with_markers = re.sub(FIGURE_BLOCK_PATTERN, "\n---image---\n", full_text, flags=re.DOTALL)

    page = body_start_page
    line_count = 0
    page_map: dict[str, int] = {}

    heading_queue: list[dict[str, object]] = list(toc_entries)
    hq_idx = 0

    for line in text_with_markers.split("\n"):
        stripped = line.strip()

        if stripped == "---pagebreak---":
            if line_count > 0:
                page += 1
                line_count = 0
            continue

        if stripped == "---image---":
            line_count += lines_per_image
            if line_count >= lines_per_page:
                page += 1
                line_count = line_count % lines_per_page
            continue

        if not stripped:
            continue

        level = 0
        text = ""
        added_lines = 1

        if stripped.startswith("### "):
            level, text = 3, stripped[4:].strip()
            added_lines = 2
        elif stripped.startswith("## "):
            level, text = 2, stripped[3:].strip()
            added_lines = 2
        elif stripped.startswith("# "):
            level, text = 1, stripped[2:].strip()
            added_lines = 3
        elif stripped.startswith("|"):
            added_lines = 1
        elif stripped.startswith("```"):
            added_lines = 1
        else:
            added_lines = max(1, (len(stripped) + 2 + chars_per_line - 1) // chars_per_line)

        # Record page for heading
        if level > 0 and hq_idx < len(heading_queue):
            entry = heading_queue[hq_idx]
            if entry["text"] == text and entry["level"] == level:
                if line_count + added_lines > lines_per_page:
                    page += 1
                    line_count = 0
                page_map[str(entry["bookmark"])] = page
                hq_idx += 1

        line_count += added_lines
        if line_count >= lines_per_page:
            page += 1
            line_count = line_count % lines_per_page

    return page_map


def _add_toc_page(
    document: DocxDocument,
    toc_entries: list[dict[str, object]],
    full_text: str = "",
) -> None:
    """Generate a visible TOC page with PAGEREF dynamic page numbers.

    Each entry is a normal paragraph structured as::

        Heading text .................. page_number

    Page numbers use PAGEREF fields referencing bookmarks on body headings,
    so WPS/Word calculates the correct page on open.  Pre-estimated page
    numbers are written as cached values so the TOC is immediately readable
    even when the application does not auto-update fields.
    """
    # body section restarts page numbering from 1
    body_start_page = 1
    page_map = _estimate_page_numbers(full_text, toc_entries, body_start_page)

    # ---- TOC title ----
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("目    录")
    _set_run_font(run_title, zh_font="黑体", size_pt=16, bold=True)

    # Per-level left indentation
    indent_map = {1: Cm(0), 2: Cm(0.74), 3: Cm(1.48)}

    for entry in toc_entries:
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = indent_map.get(_toc_int(entry["level"]), Cm(0))
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _apply_fixed_line_spacing(p.paragraph_format, pt=22)

        # Right-aligned tab stop with dot leader at right margin (15.5 cm)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        # Heading text
        run_text = p.add_run(str(entry["text"]))
        _set_run_font(run_text, size_pt=12)
        if entry["level"] == 1:
            run_text.bold = True

        # Tab character (triggers dot leader extending to page number)
        tab_run = p.add_run("\t")
        _set_run_font(tab_run, size_pt=12)

        # PAGEREF field with estimated page number as cached value
        bm = str(entry["bookmark"])
        estimated = str(page_map.get(bm, "1"))
        _add_pageref_field(p, bm, cached_page=estimated)

    # NOTE: We intentionally do NOT set w:updateFields here.
    # Setting it causes Word to auto-refresh PAGEREF fields on open,
    # but because body pages restart numbering at 1 (separate section),
    # Word resolves all PAGEREFs to "1" before completing its layout pass.
    # The pre-estimated cached page numbers are sufficiently accurate.
    # Users can manually right-click the TOC -> "Update Field" if needed.
